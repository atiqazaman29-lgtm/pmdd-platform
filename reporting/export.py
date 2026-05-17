import os
import json
from reporting.schemas import FinalResearchReport
import docx

class ExportPipeline:
    """Handles DOCX and PDF export formatting."""
    
    def export_docx(self, report: FinalResearchReport, output_path: str):
        doc = docx.Document()
        doc.add_heading(report.title, 0)
        
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(report.executive_summary)
        
        for section in report.sections:
            doc.add_heading(section.title, level=1)
            if section.ambiguity_warnings:
                p = doc.add_paragraph("WARNING: ")
                p.runs[0].bold = True
                p.add_run(" | ".join(section.ambiguity_warnings))
            
            doc.add_paragraph(section.content)
            
            if section.citations:
                doc.add_heading("Citations & Evidence", level=2)
                for cit in section.citations:
                    doc.add_paragraph(f"[{cit.segment_id}] '{cit.exact_quote}' (Reliability: {cit.reliability_score})", style='Quote')
                    
        doc.save(output_path)
        return output_path

    def export_markdown(self, report: FinalResearchReport, output_path: str):
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(f"# {report.title}\n\n")
            f.write(f"## Executive Summary\n{report.executive_summary}\n\n")
            
            for section in report.sections:
                f.write(f"## {section.title}\n")
                if section.ambiguity_warnings:
                    f.write(f"> **AMBIGUITY WARNING:** {', '.join(section.ambiguity_warnings)}\n\n")
                f.write(f"{section.content}\n\n")
        return output_path
